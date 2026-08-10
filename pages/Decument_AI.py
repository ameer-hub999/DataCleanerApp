import streamlit as st
import docx
import pypdf
import json
import os
import io
import tempfile
from pdf2docx import Converter
from groq import Groq

st.set_page_config(
    page_title="Document AI - Smart File AI", 
    page_icon="📝", 
    layout="wide"
)

# 🌐 قاموس الترجمات الكامل لجميع اللغات الخمس
UI_TEXTS = {
    "English 🇺🇸": {
        "title": "📝 Document AI & Proofreader",
        "subtitle": "Scan Word & PDF documents, get a Quality Score, and fix errors while preserving original formatting.",
        "btn_run": "🚀 Run AI Review & Analysis",
        "report_title": "📊 AI Quality Report",
        "score": "Overall Score",
        "spelling": "Spelling Errors",
        "grammar": "Grammar Issues",
        "style": "Style Suggestions",
        "breakdown": "🔍 Error Breakdown & Fixes",
        "issue_found": "Issue Found",
        "explanation": "Explanation",
        "fix": "Suggested Fix",
        "accept": "Accept Fix",
        "ignore": "Ignore",
        "no_issues": "🎉 No major errors found! Your document looks great.",
        "license_label": "🔑 Enter License / Subscription Key",
        "license_placeholder": "e.g. AMEER-PRO-2026",
        "accepted_msg": "✅ Fix Applied to Original Document (Formatting Preserved)!",
        "ignored_msg": "👁️ Issue Ignored.",
        "download_word": "📥 Download Corrected Word (.docx)",
        "download_section": "💾 Download Corrected Document",
        "invalid_key": "❌ Invalid License Key. Please enter a valid subscription key.",
        "no_key_err": "⚠️ System API key is missing. Please contact support."
    },
    "Arabic 🇸🇦": {
        "title": "📝 التدقيق والتحليل الذكي للمستندات",
        "subtitle": "افحص مستندات Word و PDF، واحصل على تقييم جودة وتصحيح للأخطاء مع الحفاظ الكامل على شكل وتنسيق الملف الأصلي.",
        "btn_run": "🚀 بدء التحليل والتدقيق الذكي",
        "report_title": "📊 تقرير جودة المستند",
        "score": "النتيجة العامة",
        "spelling": "أخطاء إملائية",
        "grammar": "أخطاء نحوية",
        "style": "تحسينات أسلوبية",
        "breakdown": "🔍 تفاصيل الأخطاء والتصحيحات",
        "issue_found": "الخطأ المكتشف",
        "explanation": "التوضيح",
        "fix": "التصحيح المقترح",
        "accept": "قبول التصحيح",
        "ignore": "تجاهل",
        "no_issues": "🎉 لم يتم العثور على أخطاء رئيسية! المستند ممتاز جداً.",
        "license_label": "🔑 أدخل كود الاشتراك / التفعيل",
        "license_placeholder": "مثال: AMEER-PRO-2026",
        "accepted_msg": "✅ تم تعديل الخطأ داخل المستند مع الحفاظ الكامل على التنسيق والخطوط!",
        "ignored_msg": "👁️ تم تجاهل الخطأ.",
        "download_word": "📥 تحميل المستند المعدّل (.docx)",
        "download_section": "💾 تحميل المستند النهائي بنفس التنسيق",
        "invalid_key": "❌ كود الاشتراك غير صحيح أو منتهي الصلاحية.",
        "no_key_err": "⚠️ مفتاح الخدمة غير مضبوط في النظام."
    },
    "Spanish 🇪🇸": {
        "title": "📝 Corrector de Documentos con IA",
        "subtitle": "Escanee documentos de Word y PDF y reciba un informe de calidad conservando el formato original.",
        "btn_run": "🚀 Ejecutar Análisis de IA",
        "report_title": "📊 Informe de Calidad de IA",
        "score": "Puntuación General",
        "spelling": "Errores Ortográficos",
        "grammar": "Problemas Gramaticales",
        "style": "Sugerencias de Estilo",
        "breakdown": "🔍 Desglose de Errores y Soluciones",
        "issue_found": "Problema Encontrado",
        "explanation": "Explicación",
        "fix": "Corrección Sugerida",
        "accept": "Aceptar",
        "ignore": "Ignorar",
        "no_issues": "🎉 ¡No se encontraron errores!",
        "license_label": "🔑 Clave de Licencia",
        "license_placeholder": "Ej. AMEER-PRO-2026",
        "accepted_msg": "✅ Corregido en el documento original",
        "ignored_msg": "👁️ Ignorado",
        "download_word": "📥 Descargar Word Corregido (.docx)",
        "download_section": "💾 Descargar Documento Final",
        "invalid_key": "❌ Licencia no válida.",
        "no_key_err": "⚠️ Clave no configurada."
    },
    "French 🇫🇷": {
        "title": "📝 AI Correcteur de Documents",
        "subtitle": "Analysez les fichiers Word et PDF et obtenez un rapport tout en conservant la mise en forme.",
        "btn_run": "🚀 Lancer l'Analyse IA",
        "report_title": "📊 Rapport de Qualité IA",
        "score": "Note Globale",
        "spelling": "Fautes d'Orthographe",
        "grammar": "Problèmes de Grammaire",
        "style": "Suggestions de Style",
        "breakdown": "🔍 Détail des Erreurs et Corrections",
        "issue_found": "Problème Détecté",
        "explanation": "Explication",
        "fix": "Correction Suggérée",
        "accept": "Accepter",
        "ignore": "Ignorer",
        "no_issues": "🎉 Aucune erreur majeure trouvée !",
        "license_label": "🔑 Clé de Licence",
        "license_placeholder": "Ex. AMEER-PRO-2026",
        "accepted_msg": "✅ Correction appliquée au document",
        "ignored_msg": "👁️ Ignoré",
        "download_word": "📥 Télécharger Word Corrigé (.docx)",
        "download_section": "💾 Télécharger le Document Final",
        "invalid_key": "❌ Clé invalide.",
        "no_key_err": "⚠️ Clé non configurée."
    },
    "German 🇩🇪": {
        "title": "📝 KI-Dokumentenprüfung",
        "subtitle": "Prüfen Sie Word- und PDF-Dokumente unter Beibehaltung der Originalformatierung.",
        "btn_run": "🚀 KI-Analyse Starten",
        "report_title": "📊 KI-Qualitätsbericht",
        "score": "Gesamtnote",
        "spelling": "Rechtschreibfehler",
        "grammar": "Grammatikfehler",
        "style": "Stilvorschläge",
        "breakdown": "🔍 Fehleraufschlüsselung & Lösungen",
        "issue_found": "Gefundener Fehler",
        "explanation": "Erklärung",
        "fix": "Vorgeschlagene Korrektur",
        "accept": "Akzeptieren",
        "ignore": "Ignorieren",
        "no_issues": "🎉 Keine Fehler gefunden!",
        "license_label": "🔑 Lizenzschlüssel",
        "license_placeholder": "z.B. AMEER-PRO-2026",
        "accepted_msg": "✅ Im Originaldokument korrigiert",
        "ignored_msg": "👁️ Ignoriert",
        "download_word": "📥 Korrigiertes Word herunterladen (.docx)",
        "download_section": "💾 Finale Datei Herunterladen",
        "invalid_key": "❌ Ungültiger Schlüssel.",
        "no_key_err": "⚠️ Schlüssel nicht konfiguriert."
    }
}

language = st.selectbox(
    "🌐 Select Document Language / اختر لغة المستند",
    ["English 🇺🇸", "Arabic 🇸🇦", "Spanish 🇪🇸", "French 🇫🇷", "German 🇩🇪"]
)

t = UI_TEXTS.get(language, UI_TEXTS["English 🇺🇸"])

st.title(t["title"])
st.write(t["subtitle"])

# 🔒 قراءة الـ API Key من Secrets أو خفية
api_key = ""
try:
    if "GROQ_API_KEY" in st.secrets:
        api_key = st.secrets["GROQ_API_KEY"]
except Exception:
    pass

if not api_key:
    api_key = os.environ.get("GROQ_API_KEY", "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")

# 🔑 أكواد الاشتراك المقبولة
valid_licenses = ["AMEER-PRO-2026", "VIP-PASS"]
try:
    if "ALLOWED_LICENSES" in st.secrets:
        valid_licenses = st.secrets["ALLOWED_LICENSES"]
except Exception:
    pass

# 🛡️ إدخال كود الاشتراك في القائمة الجانبية
st.sidebar.markdown("---")
user_license = st.sidebar.text_input(t["license_label"], placeholder=t["license_placeholder"], type="password")

uploaded_file = st.file_uploader("Upload Document (.docx, .pdf)", type=["docx", "pdf"])

def convert_pdf_to_docx_bytes(pdf_bytes):
    """تحويل PDF إلى Word مع الحفاظ على التنسيق"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(pdf_bytes)
            tmp_pdf_path = tmp_pdf.name

        tmp_docx_path = tmp_pdf_path.replace(".pdf", ".docx")

        cv = Converter(tmp_pdf_path)
        cv.convert(tmp_docx_path, start=0, end=None)
        cv.close()

        with open(tmp_docx_path, "rb") as f:
            docx_bytes = f.read()

        if os.path.exists(tmp_pdf_path):
            os.remove(tmp_pdf_path)
        if os.path.exists(tmp_docx_path):
            os.remove(tmp_docx_path)

        return docx_bytes
    except Exception:
        return None

def replace_text_preserve_formatting(paragraph, old_text, new_text):
    """استبدال الكلمات مع الحفاظ على التنسيق والـ Runs"""
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return
        
        full_text = paragraph.text.replace(old_text, new_text)
        if paragraph.runs:
            for r in paragraph.runs[1:]:
                r.text = ""
            paragraph.runs[0].text = full_text

def process_docx_replacement(doc_bytes, old_text, new_text):
    """تعديل ملف Word دون تخريب الفقرات والأنماط"""
    try:
        doc = docx.Document(io.BytesIO(doc_bytes))
        
        for p in doc.paragraphs:
            replace_text_preserve_formatting(p, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_preserve_formatting(p, old_text, new_text)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
    except Exception:
        return doc_bytes

# Session State
if "doc_analysis" not in st.session_state:
    st.session_state.doc_analysis = None

if uploaded_file is not None:
    st.success("Document uploaded successfully!")

    if st.button(t["btn_run"], type="primary"):
        if user_license.strip() not in valid_licenses:
            st.error(t["invalid_key"])
        elif not api_key:
            st.error(t["no_key_err"])
        else:
            with st.spinner("🤖 AI is analyzing your document... Please wait..."):
                try:
                    file_bytes = uploaded_file.getvalue()
                    
                    if uploaded_file.name.endswith(".pdf"):
                        converted_docx = convert_pdf_to_docx_bytes(file_bytes)
                        if converted_docx:
                            file_bytes = converted_docx

                    doc = docx.Document(io.BytesIO(file_bytes))
                    content_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                    prompt_text = f"""
                    You are an expert proofreader in language: {language}. Analyze this text for spelling, grammar, and style errors:
                    ---
                    {content_text[:3000]}
                    ---
                    Return ONLY a valid JSON object matching this schema. All text inside 'explanation' MUST be written strictly in {language}:
                    {{
                        "overall_score": 88,
                        "spelling_count": 2,
                        "grammar_count": 1,
                        "style_count": 1,
                        "issues": [
                            {{
                                "id": 0,
                                "type": "Spelling",
                                "original": "wrong_word",
                                "suggestion": "correct_word",
                                "explanation": "Explanation strictly in {language}"
                            }}
                        ]
                    }}
                    """

                    client = Groq(api_key=api_key)
                    response = client.chat.completions.create(
                        messages=[
                            {"role": "system", "content": "You are a strict JSON-only proofreading assistant."},
                            {"role": "user", "content": prompt_text}
                        ],
                        model="llama-3.3-70b-versatile",
                        response_format={"type": "json_object"}
                    )

                    parsed_data = json.loads(response.choices[0].message.content)
                    
                    issues = parsed_data.get("issues", [])
                    for idx, issue in enumerate(issues):
                        issue["id"] = idx

                    st.session_state.doc_analysis = {
                        "file_bytes": file_bytes,
                        "file_name": uploaded_file.name,
                        "text": content_text,
                        "data": parsed_data,
                        "issues": issues
                    }

                except Exception as e:
                    st.error(f"Analysis Error: {e}")

# عرض التقرير والنتائج
if st.session_state.doc_analysis is not None:
    data = st.session_state.doc_analysis["data"]
    issues = st.session_state.doc_analysis["issues"]
    
    st.markdown("---")
    st.subheader(t["report_title"])
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric(t["score"], f"{data.get('overall_score', 85)} / 100")
    col2.metric(t["spelling"], str(data.get('spelling_count', 0)))
    col3.metric(t["grammar"], str(data.get('grammar_count', 0)))
    col4.metric(t["style"], str(data.get('style_count', 0)))
    
    st.markdown("---")
    st.subheader(t["breakdown"])
    
    if not issues:
        st.balloons()
        st.success(t["no_issues"])
    else:
        for issue in list(issues):
            idx = issue["id"]
            with st.container():
                st.markdown(f"### ❌ *{t['issue_found']}:* {issue.get('original', '')} ({issue.get('type', 'Error')})")
                st.markdown(f"💡 *{t['explanation']}:* {issue.get('explanation', '')}")
                st.markdown(f"✅ *{t['fix']}:* {issue.get('suggestion', '')}")
                
                c1, c2, _ = st.columns([1, 1, 4])
                
                with c1:
                    if st.button(t["accept"], key=f"acc_{idx}"):
                        orig = issue.get('original', '')
                        sug = issue.get('suggestion', '')
                        
                        if orig and sug:
                            st.session_state.doc_analysis["text"] = st.session_state.doc_analysis["text"].replace(orig, sug)
                            
                            st.session_state.doc_analysis["file_bytes"] = process_docx_replacement(
                                st.session_state.doc_analysis["file_bytes"], orig, sug
                            )
                        
                        st.session_state.doc_analysis["issues"] = [i for i in st.session_state.doc_analysis["issues"] if i["id"] != idx]
                        st.success(t["accepted_msg"])
                        st.rerun()

                with c2:
                    if st.button(t["ignore"], key=f"ign_{idx}"):
                        st.session_state.doc_analysis["issues"] = [i for i in st.session_state.doc_analysis["issues"] if i["id"] != idx]
                        st.info(t["ignored_msg"])
                        st.rerun()
                        
                st.markdown("---")

    # 📥 زر التنزيل
    st.markdown("---")
    st.subheader(t["download_section"])
    
    download_bytes = st.session_state.doc_analysis["file_bytes"]
    out_name = f"Corrected_{st.session_state.doc_analysis['file_name'].replace('.pdf', '.docx')}"

    st.download_button(
        label=t["download_word"],
        data=download_bytes,
        file_name=out_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )